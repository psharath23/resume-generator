import os

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from resume_profiles import profiles

heading_level_2_color = "6a5380"
heading_level_1_color = "4e3d5e"
subtitle_color = "706a75"

output_folder_path = "/Users/saisharathpeddibotla/Library/Mobile Documents/com~apple~CloudDocs/Downloads"
resumes_foler_name = "resumes"

docx_generate_configs = [
    {
        "name": "SaiSharath Peddibotla",
        "email": "saisharath.peddibotla@gmail.com",
        "phone": "+13303824236",
        "linkedin": "https://www.linkedin.com/in/saisharath"
    },
    {
        "name": "SaiSharath Peddibotla",
        "email": "psharath1995@gmail.com",
        "phone": "+13303824236",
        "linkedin": "https://www.linkedin.com/in/saisharath"
    },
    {
        "name": "SaiSharath Peddibotla",
        "email": "speddibotla@gmail.com",
        "phone": "+13303824236",
        "linkedin": "https://www.linkedin.com/in/saisharath"
    }
]

def set_document_margins(section, top, bottom, left, right):
    section_properties = section._sectPr
    for margin, value in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        margin_element = OxmlElement(f'w:{margin}')
        margin_element.set(qn('w:val'), str(value))
        section_properties.append(margin_element)
        
if __name__ == "__main__":
    print(f"resumes will be output to {output_folder_path}/{resumes_foler_name}...")
    for profile in profiles:
        profile_name, specialization, summary, experiences, education, achievements, skills = profile["name"], profile["specialization"], profile["summary"], profile["experiences"], profile["education"], profile["achievements"], profile["skills"]
        
        edu_degree, edu_major, edu_from, edu_to, edu_university = education["degree"], education["major"], education["from"], education["to"], education["university"]

        for generate_config in docx_generate_configs:
            name, email, phone, linkedin = generate_config["name"], generate_config["email"], generate_config["phone"], generate_config["linkedin"]

            # Create a new Document
            doc = Document()
            
            # Add a title
            title = doc.add_paragraph(name, style='Title')
            title.alignment = WD_ALIGN_PARAGRAPH.LEFT
            title.style.font.size = Pt(26)
            title.style.font.bold = True
            title.paragraph_format.space_after = Pt(1)
            title.style.font.color.rgb = RGBColor.from_string(heading_level_1_color)
            
            # Add a Caption
            p = doc.add_paragraph(f'{phone} | {email} | {linkedin}', style='Caption')
            p.style.font.color.rgb = RGBColor.from_string(subtitle_color)
            
            # Add Subtitle
            p = doc.add_paragraph(f"{profile_name} - ({' | '.join(specialization)})", style='Subtitle')
            p.paragraph_format.space_after = Pt(0.5)
            
            # Summary
            summary_para = doc.add_paragraph(summary)
            # summary_para.paragraph_format.left_indent = Pt(10)

            # Experience Section
            p = doc.add_heading('Experience', level=1)
            p.style.font.color.rgb = RGBColor.from_string(heading_level_1_color)
            # p.style.font.size = Pt(22)

            # Add experiences (Example: Only adding one experience due to length, repeat for other experiences)
            for exp in experiences:
                exp_company, exp_pos, exp_type, exp_from, exp_to, exp_content = exp["company"], exp["position"], exp["type"], exp["from"], exp["to"], exp["content"]

                p = doc.add_heading(f"{exp_pos} - {exp_type}, {exp_company}", level=2)
                p.paragraph_format.left_indent = Pt(10)
                p.style.font.color.rgb = RGBColor.from_string(heading_level_2_color)
                p.paragraph_format.space_after = Pt(0.5)
                
                p = doc.add_heading(f'{exp_from} - {exp_to}', level=3)
                p.paragraph_format.left_indent = Pt(10)
                p.style.font.color.rgb = RGBColor.from_string(subtitle_color)
                p.paragraph_format.space_before = Pt(0.5)
                
                for c in exp_content:
                    bullet = doc.add_paragraph(c, style='List Bullet')
                    bullet.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    bullet.paragraph_format.left_indent = Inches(.5)
                    bullet.paragraph_format.space_after = Pt(0.5)

            # Education Section
            p = doc.add_heading('Education', level=1)
            p.style.font.color.rgb = RGBColor.from_string(heading_level_1_color)
            # p.style.font.size = Pt(22)

            p = doc.add_heading(f'{edu_degree}, {edu_from} - {edu_to}', level=2)
            p.style.font.color.rgb = RGBColor.from_string(heading_level_2_color)
            p.paragraph_format.left_indent = Pt(10)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            p.paragraph_format.space_after = Pt(0.5)
            
            p = doc.add_heading(f'{edu_major}', level=3)
            p.style.font.color.rgb = RGBColor.from_string(subtitle_color)
            p.paragraph_format.left_indent = Pt(10)
            p.paragraph_format.space_before = Pt(0.5)

            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            
            # p = doc.add_heading(edu_major, level=4) 
            # p.paragraph_format.left_indent = Pt(10)
            p = doc.add_paragraph(edu_university)
            p.paragraph_format.left_indent = Pt(10)
            
            # Skills Section
            p = doc.add_heading('Skills', level=1)
            p.style.font.color.rgb = RGBColor.from_string(heading_level_1_color)
            # p.style.font.size = Pt(22)
            
            for skill in skills:
                skill_area, skillset = skill["area"], skill["skillset"]
                p = doc.add_heading(f'{skill_area}', level=2)
                p.style.font.color.rgb = RGBColor.from_string(heading_level_2_color)
                p.paragraph_format.left_indent = Pt(10)
                
                p = doc.add_paragraph(skillset)
                p.paragraph_format.left_indent = Pt(10)

            # Achievements Section
            p = doc.add_heading('Achievements', level=1)
            p.style.font.color.rgb = RGBColor.from_string(heading_level_1_color)
            # p.style.font.size = Pt(22)
            
            for achievement in achievements:
                achievement_name, achievement_company, achievement_description = achievement["name"], achievement["company"], achievement["description"]
                p = doc.add_heading(f'{achievement_name} at {achievement_company}', level=2)
                p.paragraph_format.left_indent = Pt(10)
                p.style.font.color.rgb = RGBColor.from_string(heading_level_2_color)
                
                p = doc.add_paragraph(achievement_description)
                p.paragraph_format.left_indent = Pt(10)

            # Save the document
            for section in doc.sections:
                section.top_margin = Cm(2.54)
                section.bottom_margin = Cm(2.54)
                section.left_margin = Cm(2.54)
                section.right_margin = Cm(2.54)
                
            os.mkdir(f"{output_folder_path}/resumes") if resumes_foler_name not in os.listdir(output_folder_path) else print(f"{resumes_foler_name} folder exists ...")
            filename = f"{email.split('@')[0]}_{profile_name.replace(' ', '_')}_{'_'.join(specialization)}.docx"
            full_file_path = f"{output_folder_path}/{resumes_foler_name}/{filename}"
            print(f"writing resume {filename} ... ")
            doc.save(full_file_path)
        
    print("Done...!")
